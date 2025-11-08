import os
import re
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.run import Run
import PyPDF2
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
import win32com.client
import pythoncom
from pathlib import Path

from field_detector import FieldDetector
import pymupdf as fitz


class DocumentFiller:
    
    def __init__(self):
        self.field_detector = FieldDetector()
    
    def fill_document(self, template_path: str, data: Dict, 
                     output_path: str, mapping: Optional[Dict] = None) -> str:
        _, ext = os.path.splitext(template_path)
        ext = ext.lower()
        
        if ext == '.docx':
            return self._fill_docx(template_path, data, output_path, mapping)
        elif ext == '.doc':
            return self._fill_doc(template_path, data, output_path, mapping)
        elif ext == '.pdf':
            return self._fill_pdf(template_path, data, output_path, mapping)
        else:
            raise ValueError(f"Неподдерживаемый формат: {ext}")
    
    def _fill_doc(self, template_path: str, data: Dict, 
                  output_path: str, mapping: Optional[Dict] = None) -> str:
        # Конвертируем .doc в .docx с помощью pywin32
        temp_docx = template_path + 'x'  # Временный файл .docx
        
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        
        try:
            doc = word.Documents.Open(os.path.abspath(template_path))
            doc.SaveAs(os.path.abspath(temp_docx), FileFormat=16)  # 16 - формат docx
            doc.Close()
        finally:
            word.Quit()
        
        # Заполняем временный .docx
        filled_path = self._fill_docx(temp_docx, data, output_path, mapping)
        
        # Удаляем временный файл
        os.remove(temp_docx)
        
        return filled_path
    
    def _fill_docx(self, template_path: str, data: Dict, 
                   output_path: str, mapping: Optional[Dict] = None) -> str:
        doc = Document(template_path)
        
        detected_fields = self.field_detector.detect_fields_in_docx(doc)
        
        field_mappings = self.field_detector.smart_field_mapping(detected_fields, data)
        
        for para_idx, para in enumerate(doc.paragraphs):
            self._fill_paragraph(para, para_idx, field_mappings, data)
        
        for table_idx, table in enumerate(doc.tables):
            self._fill_table(table, table_idx, field_mappings, data)
        
        doc.save(output_path)
        return output_path
    
    def _fill_paragraph(self, para, para_idx: int, 
                       field_mappings: List, data: Dict):
        para_fields = [fm for fm in field_mappings 
                       if fm[0].get('location') == 'paragraph' 
                       and fm[0].get('paragraph_index') == para_idx]
        
        if not para_fields:
            return
        
        # Сортируем поля по позиции, начиная с конца, чтобы избежать сдвигов
        para_fields.sort(key=lambda x: x[0]['start'], reverse=True)
        
        for field_info, value in para_fields:
            if value is None:
                continue
                
            formatted_value = self._format_value(value, field_info)
            
            # Находим run, содержащий поле
            current_pos = 0
            for run_idx, run in enumerate(para.runs):
                run_text = run.text
                run_len = len(run_text)
                
                if current_pos <= field_info['start'] < current_pos + run_len and \
                   current_pos <= field_info['end'] <= current_pos + run_len:
                    
                    local_start = field_info['start'] - current_pos
                    local_end = field_info['end'] - current_pos
                    
                    # Заменяем текст в существующем run
                    run.text = run_text[:local_start] + formatted_value + run_text[local_end:]
                    
                    break
                
                current_pos += run_len
    
    def _fill_table(self, table, table_idx: int, 
                    field_mappings: List, data: Dict):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                self._fill_cell(cell, table_idx, row_idx, cell_idx, 
                              field_mappings, data)
    
    def _fill_cell(self, cell, table_idx: int, row_idx: int, 
                   cell_idx: int, field_mappings: List, data: Dict):
        cell_fields = [fm for fm in field_mappings 
                       if fm[0].get('location') == 'table'
                       and fm[0].get('table_index') == table_idx
                       and fm[0].get('row_index') == row_idx
                       and fm[0].get('cell_index') == cell_idx]
        
        if not cell_fields:
            return
        
        # Сортируем поля по позиции, начиная с конца
        cell_fields.sort(key=lambda x: x[0]['start'], reverse=True)
        
        for field_info, value in cell_fields:
            if value is None:
                continue
                
            formatted_value = self._format_value(value, field_info)
            
            # Предполагаем, что клетка имеет параграфы
            for para in cell.paragraphs:
                current_pos = 0
                for run_idx, run in enumerate(para.runs):
                    run_text = run.text
                    run_len = len(run_text)
                    
                    if current_pos <= field_info['start'] < current_pos + run_len and \
                       current_pos <= field_info['end'] <= current_pos + run_len:
                        
                        local_start = field_info['start'] - current_pos
                        local_end = field_info['end'] - current_pos
                        
                        # Заменяем текст в существующем run
                        run.text = run_text[:local_start] + formatted_value + run_text[local_end:]
                        
                        break
                    
                    current_pos += run_len
    
    def _fill_pdf(self, template_path: str, data: Dict, 
                  output_path: str, mapping: Optional[Dict] = None) -> str:
        detected_fields = self.field_detector.detect_fields_in_pdf(template_path)
        field_mappings = self.field_detector.smart_field_mapping(detected_fields, data)
        
        doc = fitz.open(template_path)
        
        for field_info, value in field_mappings:
            if value is None:
                continue
                
            formatted_value = self._format_value(value, field_info)
            
            page = doc[field_info['page']]
            
            # Insert text into the bbox
            bbox = field_info['bbox']
            fontsize = (bbox[3] - bbox[1]) * 0.8  # Approximate fontsize based on height
            
            page.insert_textbox(
                bbox, 
                formatted_value, 
                fontsize=fontsize, 
                fontname="helv", 
                align=0  # left align
            )
        
        doc.save(output_path, incremental=False)
        doc.close()
        
        return output_path
    
    def _format_value(self, value: Any, field_info: Dict) -> str:
        if isinstance(value, datetime):
            return value.strftime('%d.%m.%Y')
        
        if isinstance(value, (int, float)):
            field_name = field_info.get('field_name', '')
            if 'amount' in field_name or 'сумма' in field_name:
                return f"{value:,.2f}".replace(',', ' ')
            return str(value)
        
        return str(value) if value is not None else ''
    
    def fill_multiple(self, template_paths: List[str], data: Dict, 
                      output_dir: str, mapping: Optional[Dict] = None) -> List[str]:
        os.makedirs(output_dir, exist_ok=True)
        results = []
        
        for template_path in template_paths:
            base_name = os.path.basename(template_path)
            output_path = os.path.join(output_dir, f"filled_{base_name}")
            try:
                filled_path = self.fill_document(template_path, data, output_path, mapping)
                results.append(filled_path)
            except Exception as e:
                print(f"Error filling {template_path}: {e}")
        
        return results
    
    def fill_from_template_and_data(self, template_path: str, 
                                   data_source: Dict, 
                                   output_path: str,
                                   field_mapping: Optional[Dict] = None) -> Dict:
        start_time = datetime.now()
        
        try:
            result_path = self.fill_document(
                template_path, 
                data_source, 
                output_path, 
                field_mapping
            )
            
            return {
                'success': True,
                'output_path': result_path,
                'template': template_path,
                'duration': (datetime.now() - start_time).total_seconds(),
                'fields_filled': len(data_source),
                'timestamp': datetime.now().isoformat()
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'template': template_path,
                'duration': (datetime.now() - start_time).total_seconds(),
                'timestamp': datetime.now().isoformat()
            }

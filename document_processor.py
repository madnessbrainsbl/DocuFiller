import os
import re
from typing import Dict, List, Tuple, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import PyPDF2
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
import json
from datetime import datetime


class DocumentProcessor:
    
    def __init__(self):
        self.supported_formats = ['.doc', '.docx', '.pdf']
        self.field_patterns = {
            'underscore': r'_{3,}',
            'placeholder': r'\{\{(\w+)\}\}',
            'brackets': r'\[([^\]]+)\]',
            'empty_spaces': r'(?<=\s)_+(?=\s)',
        }
    
    def detect_format(self, file_path: str) -> str:
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext not in self.supported_formats:
            raise ValueError(f"Неподдерживаемый формат: {ext}")
        
        return ext
    
    def load_document(self, file_path: str) -> Any:
        doc_format = self.detect_format(file_path)
        
        if doc_format in ['.doc', '.docx']:
            return self._load_docx(file_path)
        elif doc_format == '.pdf':
            return self._load_pdf(file_path)
    
    def _load_docx(self, file_path: str) -> Document:
        try:
            return Document(file_path)
        except Exception as e:
            raise Exception(f"Ошибка загрузки DOCX: {str(e)}")
    
    def _load_pdf(self, file_path: str) -> PdfReader:
        try:
            return PdfReader(file_path)
        except Exception as e:
            raise Exception(f"Ошибка загрузки PDF: {str(e)}")
    
    def extract_text_from_docx(self, doc: Document) -> str:
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    
    def extract_text_from_pdf(self, pdf_reader: PdfReader) -> str:
        text = []
        for page in pdf_reader.pages:
            text.append(page.extract_text())
        return '\n'.join(text)
    
    def analyze_document_structure(self, file_path: str) -> Dict[str, Any]:
        doc_format = self.detect_format(file_path)
        
        if doc_format in ['.doc', '.docx']:
            return self._analyze_docx_structure(file_path)
        elif doc_format == '.pdf':
            return self._analyze_pdf_structure(file_path)
    
    def _analyze_docx_structure(self, file_path: str) -> Dict[str, Any]:
        doc = self._load_docx(file_path)
        text = self.extract_text_from_docx(doc)
        
        fields = []
        
        for pattern_name, pattern in self.field_patterns.items():
            matches = re.finditer(pattern, text)
            for match in matches:
                fields.append({
                    'type': pattern_name,
                    'position': match.start(),
                    'length': len(match.group(0)),
                    'text': match.group(0),
                    'captured': match.groups() if match.groups() else None
                })
        
        return {
            'format': '.docx',
            'text': text,
            'fields': sorted(fields, key=lambda x: x['position']),
            'paragraphs_count': len(doc.paragraphs),
            'tables_count': len(doc.tables)
        }
    
    def _analyze_pdf_structure(self, file_path: str) -> Dict[str, Any]:
        pdf_reader = self._load_pdf(file_path)
        text = self.extract_text_from_pdf(pdf_reader)
        
        fields = []
        
        for pattern_name, pattern in self.field_patterns.items():
            matches = re.finditer(pattern, text)
            for match in matches:
                fields.append({
                    'type': pattern_name,
                    'position': match.start(),
                    'length': len(match.group(0)),
                    'text': match.group(0),
                    'captured': match.groups() if match.groups() else None
                })
        
        return {
            'format': '.pdf',
            'text': text,
            'fields': sorted(fields, key=lambda x: x['position']),
            'pages_count': len(pdf_reader.pages)
        }

import json
import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from document_filler import DocumentFiller

# Load data
with open('data/example_data.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

# Create sample DOCX template
doc = Document()
doc.add_paragraph('Contract Number: {contract_number}')
doc.add_paragraph('Date: {date}')
doc.add_paragraph('Organization: {organization}')
doc.save('sample_docx_template.docx')

# Create sample PDF template
c = canvas.Canvas('sample_pdf_template.pdf', pagesize=letter)
c.drawString(100, 750, 'Contract Number: {contract_number}')
c.drawString(100, 730, 'Date: {date}')
c.drawString(100, 710, 'Organization: {organization}')
c.save()

# Create sample DOC template using pywin32
try:
    import win32com.client
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    doc = word.Documents.Add()
    doc.Content.Text = 'Contract Number: {contract_number}\r\nDate: {date}\r\nOrganization: {organization}'
    doc.SaveAs(os.path.abspath('sample_doc_template.doc'), FileFormat=0)  # 0 is wdFormatDocument
    doc.Close()
    word.Quit()
    doc_created = True
except Exception as e:
    print(f'Error creating DOC template: {e}')
    doc_created = False

# Initialize filler
filler = DocumentFiller()

# Fill DOCX
try:
    filled_docx = filler.fill_document('sample_docx_template.docx', data, 'filled_docx.docx')
    print(f'Filled DOCX saved to {filled_docx}')
except Exception as e:
    print(f'Error filling DOCX: {e}')

# Fill PDF
try:
    filled_pdf = filler.fill_document('sample_pdf_template.pdf', data, 'filled_pdf.pdf')
    print(f'Filled PDF saved to {filled_pdf}')
except Exception as e:
    print(f'Error filling PDF: {e}')

# Fill DOC if created
if doc_created:
    try:
        filled_doc = filler.fill_document('sample_doc_template.doc', data, 'filled_doc.doc')
        print(f'Filled DOC saved to {filled_doc}')
    except Exception as e:
        print(f'Error filling DOC: {e}')

# Batch filling
templates = ['sample_docx_template.docx', 'sample_pdf_template.pdf']
if doc_created:
    templates.append('sample_doc_template.doc')
output_dir = 'filled_batch'
try:
    filler.fill_multiple(templates, data, output_dir)
    print(f'Batch filling completed in {output_dir}')
except Exception as e:
    print(f'Error in batch filling: {e}')